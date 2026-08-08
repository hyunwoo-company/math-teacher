"""'문제만' 담은 시험지 DOCX 생성 (python-docx, 순수 파이썬).

풀이/변형/정답은 넣지 않는다. 각 문항의 크롭 PNG 만 번호 순서대로 삽입한다.
원본 PDF 는 수식 폰트(PUA)가 깨져 텍스트 대신 이미지를 쓴다.

이 모듈은 **블로킹**(파일 IO / 이미지 인코딩)이다.
`async def` 라우트에서 부를 때는 `run_in_threadpool` 로 감싼다.
"""

from __future__ import annotations

import io
from collections.abc import Sequence
from pathlib import Path
from typing import Final

from docx import Document
from docx.shared import Inches, Length
from PIL import Image

# 크롭 렌더 해상도(extractor.DEFAULT_DPI 와 동일). 픽셀→인치 환산 기준.
_CROP_RENDER_DPI: Final[int] = 150
# 이미지 폭 상한(A4 본문 폭에 맞춘 값). 세로 비율은 python-docx 가 유지한다.
_MAX_IMAGE_WIDTH_INCHES: Final[float] = 6.0


def _fit_width(path: Path) -> Length:
    """이미지 폭을 페이지 폭(6인치)에 맞춰 계산한다(원본보다 키우지 않음).

    python-docx 는 `width` 만 주면 세로를 원본 비율대로 맞춘다.

    Args:
        path: 크롭 PNG 경로.

    Returns:
        문서에 넣을 이미지 폭(EMU).
    """
    with Image.open(path) as image:
        width_px = image.width
    native_inches = (
        width_px / _CROP_RENDER_DPI if width_px > 0 else _MAX_IMAGE_WIDTH_INCHES
    )
    return Inches(min(native_inches, _MAX_IMAGE_WIDTH_INCHES))


def build_problems_docx(title: str, images: Sequence[tuple[int, Path]]) -> bytes:
    """상단 제목 + (번호 제목, 크롭 이미지) 목록으로 DOCX 바이트를 만든다.

    Args:
        title: 문서 상단 제목(시험지 이름).
        images: (문항 번호, 크롭 PNG 경로)를 **번호 순서대로** 담은 시퀀스.

    Returns:
        `.docx` 파일 바이트(ZIP 컨테이너, 시그니처 ``PK``).
    """
    document = Document()
    document.add_heading(title, level=0)
    for no, path in images:
        document.add_heading(f"{no}번", level=2)
        document.add_picture(str(path), width=_fit_width(path))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
