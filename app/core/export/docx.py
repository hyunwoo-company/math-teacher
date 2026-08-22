"""`ExportDoc` -> `.docx` 바이트 (python-docx, 순수 파이썬).

기존 `core/docx_export.py` 를 이식한 것이다(이미지 폭 계산 로직 동일).
원본 PDF 는 수식 폰트(PUA)가 깨져 문항은 텍스트 대신 크롭 이미지를 쓴다.

수식은 워드 네이티브 수식 개체(OMML)로 넣는다(`export/omml.py`). 분수는
가로선 위/아래로, 근호는 피근수를 덮는 선으로 워드가 직접 조판한다. 변환에
실패한 수식은 기존 평문(`MathRun.plain`)으로 폴백하고 로그를 남긴다.

이 모듈은 **블로킹**(파일 IO / 이미지 인코딩)이다.
`async def` 라우트에서 부를 때는 `run_in_threadpool` 로 감싼다.
"""

from __future__ import annotations

import io
import logging
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any, Final

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image as PilImage

from export import layout
from export.model import (
    Block,
    ExportDoc,
    Heading,
    Image,
    MathRun,
    Text,
    TextRun,
    item_spans,
)
from export.omml import UnsupportedLatexError, latex_to_omml

_LOGGER: Final[logging.Logger] = logging.getLogger(__name__)

# 크롭 렌더 해상도(extractor.DEFAULT_DPI 와 동일). 픽셀→인치 환산 기준.
_CROP_RENDER_DPI: Final[int] = 150

# 용지·여백은 `export.layout` 이 단일 소스다(hwpx 렌더러와 같은 지면을 쓴다).
# python-docx 기본 템플릿은 **Letter(8.5x11in)** 라 한국에서 인쇄하면 여백이
# 어긋나므로 여기서 A4 로 덮어쓴다.
# 이미지 폭 상한. 본문 폭을 넘으면 크롭이 여백을 침범하므로 본문 폭이 곧 상한이다.
# 세로 비율은 python-docx 가 유지한다. 예전 값 6.0in 은 Letter 본문 폭이었다.
# 1단 문서(오답노트)가 쓰는 값이다.
_MAX_IMAGE_WIDTH_INCHES: Final[float] = layout.BODY_WIDTH_MM / layout.MM_PER_INCH
# 2단 문서(시험지·변형)의 이미지 폭 상한. 이때 지면의 폭 제한은 본문 폭이 아니라
# **단 폭**(71mm)이다 — 본문 폭을 쓰면 크롭이 옆 단과 오른쪽 여백을 덮는다.
_MAX_COLUMN_IMAGE_WIDTH_INCHES: Final[float] = (
    layout.COLUMN_WIDTH_MM / layout.MM_PER_INCH
)
# `w:cols` 뒤에 올 수 있는 형제 요소들(ECMA-376 CT_SectPr 순서). 기본 템플릿에는
# `w:cols` 가 이미 있어 보통 쓰이지 않지만, 없을 때 아무 데나 끼우면 워드가 문서를
# 복구 대상으로 본다.
_COLS_SUCCESSORS: Final[tuple[str, ...]] = (
    "w:formProt",
    "w:vAlign",
    "w:noEndnote",
    "w:titlePg",
    "w:textDirection",
    "w:bidi",
    "w:rtlGutter",
    "w:docGrid",
    "w:printerSettings",
    "w:sectPrChange",
)

# 본문 폰트. python-docx 기본 템플릿은 테마 폰트(Calibri)를 쓰는데, Calibri 에는
# 위·아래첨자(ᵐ ⁿ ⁻ ₁ ₂)와 ⇒ ∘ ∠ ⋯ ≡ ✔ 글리프가 없어 평문화한 수식이 □ 로 깨진다.
# 맑은 고딕은 Windows·한글·워드에 기본 설치돼 있고 이 글자들을 모두 덮는다.
_BODY_FONT: Final[str] = "맑은 고딕"
# 본문 글자 크기. 같은 시험지를 hwpx 로 뽑으면 본문이 10pt(`charPr height="1000"`,
# 줄간격 160%)다. 글자 크기가 다르면 줄이 접히는 횟수가 달라져 두 형식의 페이지
# 수가 끝내 수렴하지 않으므로 hwpx 에 맞춘다. python-docx 기본값은 11pt 였다.
_BODY_FONT_SIZE: Final[Length] = Pt(10)
# 제목 스타일별 글자 크기(본문보다 크게, 위에서 아래로 줄어드는 사다리).
# 템플릿 기본 Title 은 26pt 로 과해서 hwpx 제목과 같은 16pt 로 낮춘다.
# Heading 1·2 는 템플릿 값(14pt·13pt)을 그대로 못박은 것이고, Heading 3 은
# 본래 Normal 을 상속해 11pt 였다 — 본문이 10pt 로 내려가도 그 크기를 유지한다.
_HEADING_FONT_SIZES: Final[Mapping[str, Length]] = {
    "Title": Pt(16),
    "Heading 1": Pt(14),
    "Heading 2": Pt(13),
    "Heading 3": Pt(11),
}
# 부속 한 줄(첫 페이지 고지 `ExportDoc.notice` / 문서 끝 출처 `ExportDoc.footer`)의
# 서식. 본문(10pt)보다 작은 회색이라 문제·해설을 읽는 데 방해가 되지 않는다.
_ASIDE_FONT_SIZE: Final[Length] = Pt(8)
_ASIDE_FONT_COLOR: Final[RGBColor] = RGBColor(0x80, 0x80, 0x80)
# 언어권별 폰트 지정. python-docx 의 `font.name` 은 ascii/hAnsi 만 건드리는데,
# 한글은 eastAsia, 수학 기호는 cs 를 따라가므로 네 가지를 모두 넣어야 한다.
_FONT_ATTRIBUTES: Final[tuple[str, ...]] = (
    "w:ascii",
    "w:eastAsia",
    "w:hAnsi",
    "w:cs",
)
# 테마 폰트 지정. 남아 있으면 워드가 위 명시값 대신 테마(Calibri)를 쓴다.
_FONT_THEME_ATTRIBUTES: Final[tuple[str, ...]] = (
    "w:asciiTheme",
    "w:eastAsiaTheme",
    "w:hAnsiTheme",
    "w:cstheme",
)


def _apply_font(style: Any, name: str, size: Length) -> None:
    """스타일에 폰트와 글자 크기를 못박는다(테마 해제 + 복합 스크립트 동기화).

    python-docx 의 `font.name` 은 `w:ascii` 와 `w:hAnsi` 만 쓰고 테마 속성
    (`w:asciiTheme` 등)은 그대로 둔다. 테마 속성이 남아 있으면 워드가 그쪽을
    우선해 Calibri 로 되돌아가므로 XML 에서 직접 지우고 네 언어권을 다 채운다.
    `font.size` 도 `w:sz` 만 쓰므로 `w:szCs` 를 같은 값으로 맞춰 준다.

    Args:
        style: python-docx 스타일 객체(`ParagraphStyle`). 정확한 타입이
            `Styles.__getitem__` 반환형(`BaseStyle`)으로 좁혀지지 않아 `Any` 다.
        name: 지정할 폰트 이름.
        size: 지정할 글자 크기.
    """
    style.font.name = name
    style.font.size = size
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attribute in _FONT_ATTRIBUTES:
        fonts.set(qn(attribute), name)
    for attribute in _FONT_THEME_ATTRIBUTES:
        fonts.attrib.pop(qn(attribute), None)
    complex_size = rpr.find(qn("w:szCs"))
    if complex_size is None:
        complex_size = OxmlElement("w:szCs")
        rpr.get_or_add_sz().addnext(complex_size)
    complex_size.set(qn("w:val"), str(int(size.pt * 2)))


def _tighten(document: DocxDocument) -> None:
    """문단 여백·줄간격·글자 크기를 hwpx 쪽에 맞춰 페이지 수를 정상화한다.

    python-docx 기본 템플릿의 `docDefaults` 는 문단마다 뒤 여백 10pt
    (`w:after="200"`)와 줄간격 1.15배(`w:line="276"`)를 준다. 이 렌더러는 평문을
    줄 단위로 문단화하므로 그 여백이 문단 수만큼 곱해져, 같은 내용이 hwpx 14쪽
    대비 docx 74쪽으로 불어났다(설계 §2-1 실측).

    글자 크기도 같이 맞춘다. hwpx 는 본문 10pt 인데 docx 기본은 11pt 라, 여백을
    없애도 긴 줄이 더 자주 접혀 페이지 수가 수렴하지 않는다.

    Args:
        document: 방금 만든 빈 문서. 본문을 채우기 전에 부른다.
    """
    normal = document.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    # SINGLE 은 `w:line="240" w:lineRule="auto"`(=1.0배)를 만든다.
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _apply_font(normal, _BODY_FONT, _BODY_FONT_SIZE)

    for name, size in _HEADING_FONT_SIZES.items():
        try:
            style = document.styles[name]
        except KeyError:
            continue
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(2)
        _apply_font(style, _BODY_FONT, size)


def _set_page(document: DocxDocument) -> None:
    """용지를 A4 로, 여백을 hwpx 쪽 실측값에 맞춘다.

    python-docx 기본 `sectPr` 은 Letter(8.5x11in)라 한국에서 인쇄하면 여백이
    어긋난다. 같은 문서를 hwpx 로 뽑으면 A4 로 나가므로 두 형식이 애초에 다른
    용지였고, 페이지 수를 비교하는 것 자체가 성립하지 않았다.

    Args:
        document: 방금 만든 빈 문서. 본문을 채우기 전에 부른다.
    """
    section = document.sections[0]
    section.page_width = Mm(layout.PAGE_WIDTH_MM)
    section.page_height = Mm(layout.PAGE_HEIGHT_MM)
    section.left_margin = Mm(layout.MARGIN_SIDE_MM)
    section.right_margin = Mm(layout.MARGIN_SIDE_MM)
    section.top_margin = Mm(layout.MARGIN_TOP_MM)
    section.bottom_margin = Mm(layout.MARGIN_BOTTOM_MM)


def _set_columns(document: DocxDocument) -> None:
    """구역을 좌우 2단으로 만들고 단 사이에 세로 구분선을 세운다.

    워드는 단 설정을 구역 속성(`w:sectPr/w:cols`)으로 갖는다 — 문단마다 거는
    것이 아니라 구역 전체에 걸린다. python-docx 에 고수준 API 가 없어 요소를
    직접 만진다(기본 템플릿에는 `<w:cols w:space="720"/>` 가 이미 있다).

    `w:sep="1"` 이 **중앙 세로선**이고, `w:space` 는 단 사이 간격(twip)이다.
    제목과 고지도 같은 구역에 있으므로 함께 2단 안에 들어간다(hwpx 렌더러와
    같은 모양이 되도록 일부러 구역을 나누지 않는다).

    Args:
        document: 방금 만든 빈 문서.
    """
    sect_pr = document.sections[0]._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.insert_element_before(cols, *_COLS_SUCCESSORS)
    cols.set(qn("w:num"), str(layout.COLUMN_COUNT))
    cols.set(qn("w:space"), str(Mm(layout.COLUMN_GAP_MM).twips))
    cols.set(qn("w:sep"), "1")


def _fit_width(path: Path, max_inches: float) -> Length:
    """이미지 폭을 지면이 허용하는 폭에 맞춘다(원본보다 키우지 않음).

    python-docx 는 `width` 만 주면 세로를 원본 비율대로 맞춘다.

    Args:
        path: 크롭 PNG 경로.
        max_inches: 폭 상한(인치). 1단은 본문 폭(150mm = 5.91in), 2단은 단 폭
            (71mm = 2.80in)이다.

    Returns:
        문서에 넣을 이미지 폭(EMU).
    """
    with PilImage.open(path) as image:
        width_px = image.width
    native_inches = width_px / _CROP_RENDER_DPI if width_px > 0 else max_inches
    return Inches(min(native_inches, max_inches))


def _add_math(paragraph: Paragraph, run: MathRun) -> None:
    """문단 끝에 워드 수식 개체를 붙인다. 변환 실패 시 평문으로 폴백한다.

    `m:oMath` 는 `w:p` 의 정식 자식이라(ECMA-376 Part 1 §22.1.2 / EG_PContent)
    문단 요소에 그대로 덧붙일 수 있다.

    Args:
        paragraph: 붙일 문단.
        run: 수식 런.
    """
    try:
        math_xml = latex_to_omml(run.latex)
    except UnsupportedLatexError as error:
        _LOGGER.info("OMML 변환 실패, 평문으로 폴백: %r (%s)", run.latex, error)
        paragraph.add_run(run.plain)
        return
    paragraph._p.append(parse_xml(math_xml))


def _add_aside(document: DocxDocument, text: str) -> None:
    """작은 회색 한 줄을 문단으로 붙인다(첫 페이지 고지 / 문서 끝 출처 공용).

    Args:
        document: 대상 문서.
        text: 넣을 한 줄.
    """
    run = document.add_paragraph().add_run(text)
    run.font.size = _ASIDE_FONT_SIZE
    run.font.color.rgb = _ASIDE_FONT_COLOR


def _add_text(document: DocxDocument, block: Text) -> list[Paragraph]:
    """본문 블록을 문단들로 넣는다.

    한 문단에 개행을 그대로 넣으면 워드가 줄바꿈으로 표시하지 않으므로 줄마다
    문단을 만든다.

    Args:
        document: 대상 문서.
        block: 본문 블록.

    Returns:
        만든 문단들(호출자가 문항 단위 조판 속성을 걸 수 있게 돌려준다).
    """
    if block.lines is None:
        # 수식이 없는 블록. 예전 경로 그대로다.
        return [document.add_paragraph(line) for line in block.text.split("\n")]
    paragraphs: list[Paragraph] = []
    for runs in block.lines:
        paragraph = document.add_paragraph()
        paragraphs.append(paragraph)
        for run in runs:
            if isinstance(run, TextRun):
                paragraph.add_run(run.text)
            else:
                _add_math(paragraph, run)
    return paragraphs


def _add_block(
    document: DocxDocument, block: Block, *, max_image_inches: float
) -> list[Paragraph]:
    """블록 하나를 렌더하고 그때 만든 문단들을 돌려준다.

    문단을 돌려주는 이유는 문항 단위로 조판 속성을 걸어야 하기 때문이다
    (`_apply_item_layout`). 문서를 다 만든 뒤 `document.paragraphs` 를 훑어
    되찾는 방법도 있지만, 434문항 문서에서는 문항마다 전체 문단 목록을 새로
    만들어 O(문항수 x 문단수) 가 된다.

    이미지는 `Document.add_picture` 와 **같은 XML** 을 만든다(그 메서드가
    `add_paragraph().add_run()` 뒤에 `run.add_picture` 를 부르는 것과 같다).
    문단 객체를 직접 잡으려고 두 줄로 편 것뿐이다.

    Args:
        document: 대상 문서.
        block: 렌더할 블록.
        max_image_inches: 이미지 폭 상한(인치).

    Returns:
        만든 문단들. 페이지 나눔은 문항 구간 밖이므로 빈 목록을 준다.
    """
    if isinstance(block, Heading):
        return [document.add_heading(block.text, level=block.level)]
    if isinstance(block, Image):
        paragraph = document.add_paragraph()
        paragraph.add_run().add_picture(
            str(block.path), width=_fit_width(block.path, max_image_inches)
        )
        return [paragraph]
    if isinstance(block, Text):
        return _add_text(document, block)
    # 남은 것은 `PageBreak` 뿐이다(`Block` 유니온 4종 중 셋을 위에서 처리했다).
    # `w:br w:type="page"` 를 담은 문단 하나를 만든다(python-docx 내장).
    # python-docx 가 이 메서드에만 주석을 달지 않아 strict mypy 가 막는다.
    document.add_page_break()  # type: ignore[no-untyped-call]
    return []


def _apply_item_layout(paragraphs: Sequence[Paragraph]) -> None:
    """문항 하나를 이루는 문단들이 단·페이지를 넘어 쪼개지지 않게 한다.

    높이를 우리가 계산하지 않는다. 크롭 높이는 알아도 텍스트가 몇 줄로 접힐지는
    폰트·단 폭에 따라 달라져 예측이 어긋나고, 워드와 한글이 다르게 조판한다.
    대신 조판 엔진에 "이 문단들은 쪼개지 마라"(`w:keepLines`)와 "다음 문단과
    붙어라"(`w:keepNext`)를 걸어 두면, 남은 공간에 안 들어갈 때 워드가 문항을
    **통째로** 다음 단이나 다음 페이지로 넘긴다.

    마지막 문단만 `keepNext` 를 끈다. 켜 두면 다음 문항까지 끌려와 문서 전체가
    한 덩이가 된다. 워드 기본 제목 스타일이 `keepNext` 를 켜 두는 경우가 있어
    상속에 맡기지 않고 **명시적으로 끈다**.

    첫 문단(= 문항 제목)에는 문항 간 간격을 준다. 문항 안 문단 사이는 그대로
    촘촘하게 둔다 — 그래야 한 문항이 한 덩이로 보인다.

    **한계**: 문항 하나가 한 단 높이보다 크면 어떤 설정으로도 넘길 곳이 없다.
    그런 문항은 결국 단 사이에서 쪼개지거나 한 단을 통째로 차지한다.

    Args:
        paragraphs: 한 문항이 만든 문단들(제목부터 순서대로). 비면 아무 일도 없다.
    """
    if not paragraphs:
        return
    last = len(paragraphs) - 1
    for index, paragraph in enumerate(paragraphs):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True
        paragraph_format.keep_with_next = index != last
    paragraphs[0].paragraph_format.space_before = Pt(layout.ITEM_GAP_PT)


def build_docx(doc: ExportDoc) -> bytes:
    """`ExportDoc` 을 `.docx` 바이트로 렌더한다.

    본문 텍스트는 줄 단위로 문단을 나눈다. 한 문단에 개행을 그대로 넣으면
    워드가 줄바꿈으로 표시하지 않기 때문이다. 그래서 문단이 아주 많아지므로
    `_tighten` 으로 문단 여백을 0 으로 눌러 두고, `_set_page` 로 용지를 A4 로
    맞춘 다음 시작한다.

    `doc.notice` 가 있으면 제목 바로 아래(첫 페이지)에, `doc.footer` 는 문서 맨
    끝에 작은 회색 한 줄로 넣는다. 둘 다 없으면 예전과 똑같은 문서다.

    `doc.two_column` 이면 구역을 좌우 2단(중앙 세로선)으로 세우고, 문항마다
    쪼개짐 방지와 문항 간 간격을 건다(`_set_columns` / `_apply_item_layout`).
    1단 문서(오답노트)는 예전과 똑같이 나간다 — 문단 속성도 붙지 않는다.

    Args:
        doc: 렌더할 문서.

    Returns:
        `.docx` 파일 바이트(ZIP 컨테이너, 시그니처 ``PK``).
    """
    document = Document()
    _set_page(document)
    if doc.two_column:
        _set_columns(document)
    _tighten(document)
    max_image_inches = (
        _MAX_COLUMN_IMAGE_WIDTH_INCHES if doc.two_column else _MAX_IMAGE_WIDTH_INCHES
    )
    document.add_heading(doc.title, level=0)
    if doc.notice:
        # 고지는 **제목 바로 아래**(= 첫 페이지)다. 읽기 전에 보여야 의미가 있다.
        _add_aside(document, doc.notice)
    blocks = list(doc.blocks)
    # 1단 문서는 문항 구간을 찾지 않는다 — 걸 속성이 없으므로 예전 경로 그대로다.
    spans = item_spans(blocks) if doc.two_column else {}
    index = 0
    while index < len(blocks):
        stop = spans.get(index)
        if stop is None:
            _add_block(document, blocks[index], max_image_inches=max_image_inches)
            index += 1
            continue
        _apply_item_layout(
            [
                paragraph
                for block in blocks[index:stop]
                for paragraph in _add_block(
                    document, block, max_image_inches=max_image_inches
                )
            ]
        )
        index = stop
    if doc.footer:
        # 출처는 문서 맨 끝 한 줄. 본문과 섞이지 않게 작은 회색 글씨로 낸다.
        _add_aside(document, doc.footer)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


__all__ = ["build_docx"]
